# tools/translate_pdf_range.py
"""
PDF 範圍翻譯工具

從 PDF 指定範圍直接翻譯並插入 Word 模板

【起點】從首次出現 "OVERVIEW OF ENERGY SOURCES AND SAFEGUARDS" 的頁面開始（包含該頁）
【終點】在文件章節標題進入 "ATTACHMENT TO TEST REPORT" 時停止（不包含該頁）
       若整份 PDF 中未出現上述章節標題，則翻譯至 PDF 最後一頁為止

【限制】
- 不得改變表格結構、欄位或順序
- 不得摘要、重寫或補充原文
- 僅做逐句、逐表格的忠實翻譯
"""

import os
import sys
import json
import re
import argparse
from pathlib import Path
from typing import List, Tuple, Optional, Dict
from copy import deepcopy

import pdfplumber
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 載入環境變數
from dotenv import load_dotenv
load_dotenv()

# 添加專案根目錄
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from core.llm_translator import get_translator


# ============================================================
# 常數定義
# ============================================================
START_MARKER = "OVERVIEW OF ENERGY SOURCES AND SAFEGUARDS"
END_MARKERS = [
    "ATTACHMENT TO TEST REPORT",
    "ATTACHMENTS TO TEST REPORT",
]


def find_translation_range(pdf_path: str) -> Tuple[int, int]:
    """
    找出 PDF 翻譯範圍

    Returns:
        (start_page, end_page): 0-indexed 頁碼範圍 [start, end)
    """
    with pdfplumber.open(pdf_path) as pdf:
        start_page = None
        end_page = len(pdf.pages)  # 預設到最後一頁

        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""

            # 找起點
            if start_page is None and START_MARKER in text:
                start_page = i
                print(f"[翻譯範圍] 起點: Page {i + 1} (找到 '{START_MARKER}')")

            # 找終點
            for end_marker in END_MARKERS:
                if end_marker in text:
                    end_page = i  # 不包含此頁
                    print(f"[翻譯範圍] 終點: Page {i + 1} (找到 '{end_marker}')")
                    break

            if end_page != len(pdf.pages):
                break

        if start_page is None:
            # 如果找不到起點標題，從第一頁開始
            start_page = 0
            print(f"[翻譯範圍] 警告：未找到起點標題，從 Page 1 開始")

        if end_page == len(pdf.pages):
            print(f"[翻譯範圍] 終點: Page {end_page} (PDF 最後一頁)")

        print(f"[翻譯範圍] 共 {end_page - start_page} 頁 (Page {start_page + 1} ~ {end_page})")

        return start_page, end_page


def extract_tables_from_range(pdf_path: str, start_page: int, end_page: int) -> List[Dict]:
    """
    從 PDF 指定範圍抽取所有表格，包含合併儲存格和背景色資訊

    完全保留 PDF 原始格式：
    - 欄位數量：按 PDF 原有結構
    - 合併儲存格：分析空白欄位推斷
    - 背景色：從 PDF 矩形物件讀取

    Returns:
        list of dict: [
            {
                'page': 9,
                'rows': [[cell1, cell2, ...], ...],
                'col_count': 4,
                'merge_info': [  # 合併儲存格資訊
                    {'row': 0, 'col': 0, 'colspan': 5},
                    ...
                ],
                'row_backgrounds': [True, False, True, ...]  # 每行是否有背景色
            }, ...
        ]
    """
    tables = []

    with pdfplumber.open(pdf_path) as pdf:
        for page_idx in range(start_page, end_page):
            page = pdf.pages[page_idx]
            page_num = page_idx + 1  # 1-indexed for display

            # 抽取頁面上所有灰色背景矩形
            filled_rects = _extract_filled_rects(page)

            try:
                # 使用 find_tables 來取得表格物件（包含 cells 位置資訊）
                page_table_objs = page.find_tables({
                    'vertical_strategy': 'lines',
                    'horizontal_strategy': 'lines',
                    'intersection_tolerance': 3,
                    'snap_tolerance': 3,
                    'join_tolerance': 3,
                })
            except Exception as e:
                print(f"[警告] Page {page_num} 表格抽取失敗: {e}")
                continue

            for table_obj in page_table_objs:
                if not table_obj:
                    continue

                # 抽取表格資料
                tbl = table_obj.extract()
                if not tbl:
                    continue

                # 正規化表格資料
                rows = []
                max_cols = 0
                for row in tbl:
                    if row:
                        normalized_row = [_normalize_cell(c) for c in row]
                        rows.append(normalized_row)
                        max_cols = max(max_cols, len(normalized_row))

                if not rows:
                    continue

                # 過濾掉 PDF 頁眉表格（通常只有 1-2 行且包含特定關鍵字）
                first_row_text = ' '.join(rows[0]) if rows else ''
                is_header_table = (
                    len(rows) <= 2 and
                    ('IEC 62368-1' in first_row_text or
                     'Requirement + Test' in first_row_text or
                     'Result - Remark' in first_row_text or
                     'Clause' in first_row_text and 'Verdict' in first_row_text)
                )

                if is_header_table:
                    continue

                # 分析合併儲存格
                merge_info = _analyze_merged_cells(table_obj, rows, max_cols)

                # 分析每行的背景色
                row_backgrounds = _analyze_row_backgrounds(table_obj, rows, max_cols, filled_rects)

                tables.append({
                    'page': page_num,
                    'rows': rows,
                    'col_count': max_cols,
                    'merge_info': merge_info,
                    'row_backgrounds': row_backgrounds
                })

    print(f"[抽取] 共抽取 {len(tables)} 個表格")
    return tables


def _extract_filled_rects(page) -> List[Dict]:
    """
    抽取頁面上所有灰色背景矩形

    Returns:
        list of dict: [{'x0': ..., 'top': ..., 'x1': ..., 'bottom': ..., 'color': 0.898}, ...]
    """
    filled_rects = []

    for rect in page.rects:
        # 檢查是否為填充矩形
        if not rect.get('fill'):
            continue

        color = rect.get('non_stroking_color')
        if color is None:
            continue

        # 只保留灰色背景 (約 0.7-0.99，排除白色 1.0 和黑色 0.0)
        if isinstance(color, (int, float)) and 0.7 < color < 1.0:
            filled_rects.append({
                'x0': rect['x0'],
                'top': rect['top'],
                'x1': rect['x1'],
                'bottom': rect['bottom'],
                'color': color,
            })

    return filled_rects


def _analyze_row_backgrounds(table_obj, rows: List[List[str]], col_count: int, filled_rects: List[Dict]) -> List[bool]:
    """
    分析表格每行是否有背景色

    透過檢查每行第一個 cell 是否被灰色矩形覆蓋來判斷

    Returns:
        list of bool: [True, False, True, ...] 每行是否有背景色
    """
    row_backgrounds = []
    cells = table_obj.cells  # 每個 cell 的座標 (x0, top, x1, bottom)
    tolerance = 5  # 座標容差

    for r_idx in range(len(rows)):
        # 找到此行第一個 cell 的座標
        cell_idx = r_idx * col_count
        if cell_idx >= len(cells):
            row_backgrounds.append(False)
            continue

        cx0, ctop, cx1, cbottom = cells[cell_idx]

        # 檢查是否有灰色矩形覆蓋此 cell
        has_background = False
        for rect in filled_rects:
            if (rect['x0'] <= cx0 + tolerance and
                rect['x1'] >= cx1 - tolerance and
                rect['top'] <= ctop + tolerance and
                rect['bottom'] >= cbottom - tolerance):
                has_background = True
                break

        row_backgrounds.append(has_background)

    return row_backgrounds


def _analyze_merged_cells(table_obj, rows: List[List[str]], col_count: int) -> List[Dict]:
    """
    分析表格的合併儲存格 - 基於 PDF cell 座標精確計算

    透過分析 pdfplumber 的 cell 座標來判斷每個 cell 跨越了幾欄幾行

    Returns:
        list of dict: [
            {'row': 0, 'col': 0, 'colspan': 5, 'rowspan': 1},
            {'row': 3, 'col': 2, 'colspan': 3, 'rowspan': 2},
            ...
        ]
    """
    merge_info = []

    if not rows or col_count == 0:
        return merge_info

    cells = table_obj.cells
    if not cells:
        return merge_info

    # 找出所有 X 和 Y 座標邊界
    x_coords = sorted(set(round(c[0], 0) for c in cells) | set(round(c[2], 0) for c in cells))
    y_coords = sorted(set(round(c[1], 0) for c in cells) | set(round(c[3], 0) for c in cells))

    if len(x_coords) < 2 or len(y_coords) < 2:
        return merge_info

    # 分析每個 cell 的合併情況
    for cell in cells:
        x0, top, x1, bottom = cell

        try:
            start_col = x_coords.index(round(x0, 0))
            end_col = x_coords.index(round(x1, 0))
            start_row = y_coords.index(round(top, 0))
            end_row = y_coords.index(round(bottom, 0))
        except ValueError:
            continue

        colspan = end_col - start_col
        rowspan = end_row - start_row

        # 只記錄有合併的 cell（colspan > 1 或 rowspan > 1）
        if colspan > 1 or rowspan > 1:
            # 確保 row 在 rows 範圍內
            if start_row < len(rows):
                merge_info.append({
                    'row': start_row,
                    'col': start_col,
                    'colspan': colspan,
                    'rowspan': rowspan
                })

    return merge_info


def _normalize_cell(cell) -> str:
    """正規化儲存格內容"""
    if cell is None:
        return ""
    text = str(cell)
    # 移除多餘空白
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def translate_tables(tables: List[Dict]) -> List[Dict]:
    """
    翻譯所有表格內容

    保持表格結構不變，僅翻譯儲存格內容
    """
    translator = get_translator()
    translated_tables = []

    # 收集所有需要翻譯的文本
    all_texts = []
    text_positions = []  # (table_idx, row_idx, col_idx)

    for t_idx, table in enumerate(tables):
        for r_idx, row in enumerate(table['rows']):
            for c_idx, cell in enumerate(row):
                if cell and _needs_translation(cell):
                    all_texts.append(cell)
                    text_positions.append((t_idx, r_idx, c_idx))

    print(f"[翻譯] 共 {len(all_texts)} 個儲存格需要翻譯")

    # 批次翻譯
    if all_texts:
        translated_texts = translator.translate_batch(all_texts)
    else:
        translated_texts = []

    # 建立翻譯結果的深拷貝
    translated_tables = deepcopy(tables)

    # 將翻譯結果填回表格
    for i, (t_idx, r_idx, c_idx) in enumerate(text_positions):
        translated_tables[t_idx]['rows'][r_idx][c_idx] = translated_texts[i]

    return translated_tables


def _set_table_borders(table):
    """手動設定表格框線"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _set_cell_shading(cell, color: str):
    """
    設定儲存格背景色

    Args:
        cell: Word 儲存格物件
        color: 16 進位顏色碼 (如 "D9D9D9")
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 移除現有的 shading
    existing_shd = tcPr.find(qn('w:shd'))
    if existing_shd is not None:
        tcPr.remove(existing_shd)

    # 建立新的 shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def _apply_merge_to_table(table, merge_info: List[Dict], merged_cells: set = None):
    """
    手動設定表格的合併儲存格（正確處理垂直合併）

    python-docx 的 merge() 方法有 bug，垂直合併時會錯誤地設定 vMerge
    這個函數直接操作 XML 來正確處理

    Word 合併邏輯（colspan=3, rowspan=2 的例子）：
    - 第一行 cell: gridSpan=3, vMerge=restart
    - 第二行 cell: gridSpan=3, vMerge（無 val = 繼續）
    - 每行只需要設定起始 column 的 cell，其餘被 gridSpan 覆蓋

    Args:
        table: Word 表格物件
        merge_info: [{'row': 0, 'col': 0, 'colspan': 5, 'rowspan': 1}, ...]
        merged_cells: set of (row, col) tuples that are covered by merges
    """
    if merged_cells is None:
        merged_cells = set()

    # 直接從 XML 取得所有 tr (行) 元素
    tbl = table._tbl
    tr_list = tbl.findall(qn('w:tr'))

    for m in merge_info:
        r_idx = m['row']
        c_idx = m['col']
        colspan = m.get('colspan', 1)
        rowspan = m.get('rowspan', 1)

        # 處理每個受影響的行
        for dr in range(rowspan):
            row_idx = r_idx + dr
            if row_idx >= len(tr_list):
                continue

            tr = tr_list[row_idx]
            tc_list = tr.findall(qn('w:tc'))

            # 只處理起始 column，水平合併的其他 column 不需要特別處理
            col_idx = c_idx
            if col_idx >= len(tc_list):
                continue

            tc = tc_list[col_idx]

            # 如果是被覆蓋的 cell（非起始行），清空內容
            if (row_idx, col_idx) in merged_cells:
                for p in tc.findall(qn('w:p')):
                    for r in list(p):
                        if r.tag != qn('w:pPr'):
                            p.remove(r)

            # 取得或建立 tcPr
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # 設定 gridSpan（水平合併）- 每一行都需要設定
            if colspan > 1:
                grid_span = tcPr.find(qn('w:gridSpan'))
                if grid_span is None:
                    grid_span = OxmlElement('w:gridSpan')
                    tcPr.append(grid_span)
                grid_span.set(qn('w:val'), str(colspan))

            # 設定 vMerge（垂直合併）
            if rowspan > 1:
                # 移除現有的 vMerge
                existing_vmerge = tcPr.find(qn('w:vMerge'))
                if existing_vmerge is not None:
                    tcPr.remove(existing_vmerge)

                # 建立新的 vMerge
                v_merge = OxmlElement('w:vMerge')
                if dr == 0:
                    # 第一行：restart（開始合併）
                    v_merge.set(qn('w:val'), 'restart')
                # 其他行：不設定 val 屬性（繼續合併）
                tcPr.append(v_merge)


def _needs_translation(text: str) -> bool:
    """判斷文本是否需要翻譯（包含英文）"""
    if not text:
        return False

    # 如果已經是純中文，不需要翻譯
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    total_chars = len(re.findall(r'[a-zA-Z\u4e00-\u9fff]', text))

    if total_chars == 0:
        return False

    # 如果中文比例超過 90%，不需要翻譯
    if chinese_chars / total_chars > 0.9:
        return False

    return True


def insert_tables_to_template(
    template_path: str,
    translated_tables: List[Dict],
    output_path: str,
    insert_after_table_idx: int = 3  # 在表格 3 (試驗樣品特性) 之後插入
):
    """
    將翻譯後的表格插入模板

    Args:
        template_path: 模板路徑
        translated_tables: 翻譯後的表格列表
        output_path: 輸出路徑
        insert_after_table_idx: 在第幾個表格之後插入 (0-indexed)
    """
    doc = Document(template_path)

    print(f"[插入] 模板共有 {len(doc.tables)} 個表格")
    print(f"[插入] 將在表格 {insert_after_table_idx} 之後插入 {len(translated_tables)} 個新表格")

    # 找到插入位置（表格 3 之後的段落）
    if insert_after_table_idx < len(doc.tables):
        last_table = doc.tables[insert_after_table_idx]
        # 在表格後面插入新內容
        insert_element = last_table._tbl

        # 在 Table 3 之後插入分頁符，讓翻譯內容從第 5 頁開始
        from docx.oxml import OxmlElement
        page_break_para = OxmlElement('w:p')
        page_break_run = OxmlElement('w:r')
        page_break = OxmlElement('w:br')
        page_break.set(qn('w:type'), 'page')
        page_break_run.append(page_break)
        page_break_para.append(page_break_run)
        insert_element.addnext(page_break_para)
        insert_element = page_break_para
    else:
        # 如果表格不夠，在文件末尾插入
        insert_element = doc.element.body[-1]

    # 逐個插入表格
    for t_idx, table_data in enumerate(translated_tables):
        rows = table_data['rows']
        col_count = table_data['col_count']
        merge_info = table_data.get('merge_info', [])
        row_backgrounds = table_data.get('row_backgrounds', [])

        if not rows:
            continue

        # 建立新表格
        new_table = doc.add_table(rows=len(rows), cols=col_count)
        # 嘗試設定表格樣式，如果失敗則跳過
        try:
            new_table.style = 'Table Grid'
        except KeyError:
            # 樣式不存在，手動設定框線
            _set_table_borders(new_table)

        # 建立合併查詢表（用於跳過已被合併的 cell）
        merged_cells = set()  # (row, col) 已被合併覆蓋的 cell
        for m in merge_info:
            r = m['row']
            c = m['col']
            colspan = m.get('colspan', 1)
            rowspan = m.get('rowspan', 1)
            # 記錄被合併覆蓋的所有 cell（排除起始 cell）
            for dr in range(rowspan):
                for dc in range(colspan):
                    if dr > 0 or dc > 0:
                        merged_cells.add((r + dr, c + dc))

        # 先填入資料（在合併前），使用直接 XML 存取
        tbl = new_table._tbl
        tr_list = tbl.findall(qn('w:tr'))

        for r_idx, row in enumerate(rows):
            if r_idx >= len(tr_list):
                continue

            tr = tr_list[r_idx]
            tc_list = tr.findall(qn('w:tc'))

            # 判斷此行是否需要灰色背景
            needs_gray_bg = row_backgrounds[r_idx] if r_idx < len(row_backgrounds) else False

            for c_idx, cell_text in enumerate(row):
                if c_idx >= len(tc_list):
                    continue

                # 使用 python-docx 的 _Cell 包裝來設定文字和格式
                from docx.table import _Cell
                tc = tc_list[c_idx]
                cell = _Cell(tc, new_table)
                cell.text = cell_text or ""

                # 設定字型
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = '標楷體'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

                # 套用灰色背景
                if needs_gray_bg:
                    _set_cell_shading(cell, "D9D9D9")

        # 最後才套用合併（避免影響資料填入）
        _apply_merge_to_table(new_table, merge_info, merged_cells)

        # 移動表格到正確位置
        insert_element.addnext(new_table._tbl)
        insert_element = new_table._tbl

        if (t_idx + 1) % 10 == 0:
            print(f"  已插入 {t_idx + 1}/{len(translated_tables)} 個表格...")

    # 儲存
    doc.save(output_path)
    print(f"[完成] 輸出: {output_path}")


def process_pdf_to_docx(
    pdf_path: str,
    template_path: str,
    output_path: str
):
    """
    主流程：PDF → 翻譯 → DOCX

    Args:
        pdf_path: CB PDF 路徑
        template_path: CNS 模板路徑
        output_path: 輸出 DOCX 路徑
    """
    print("=" * 60)
    print("PDF 範圍翻譯工具")
    print("=" * 60)

    # Step 1: 找出翻譯範圍
    print("\n[Step 1] 識別翻譯範圍...")
    start_page, end_page = find_translation_range(pdf_path)

    # Step 2: 抽取表格
    print("\n[Step 2] 抽取表格...")
    tables = extract_tables_from_range(pdf_path, start_page, end_page)

    # Step 3: 翻譯表格
    print("\n[Step 3] 翻譯表格...")
    translated_tables = translate_tables(tables)

    # Step 4: 插入模板
    print("\n[Step 4] 插入模板...")
    insert_tables_to_template(template_path, translated_tables, output_path)

    print("\n" + "=" * 60)
    print("處理完成！")
    print("=" * 60)


def main():
    parser = argparse.ArgumentParser(description='PDF 範圍翻譯工具')
    parser.add_argument('--pdf', required=True, help='CB PDF 路徑')
    parser.add_argument('--template', default='templates/CNS_15598_1_109_template_clean.docx',
                        help='CNS 模板路徑')
    parser.add_argument('--out', required=True, help='輸出 DOCX 路徑')

    args = parser.parse_args()

    process_pdf_to_docx(args.pdf, args.template, args.out)


if __name__ == "__main__":
    main()
